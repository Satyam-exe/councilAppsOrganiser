import io
import os
import subprocess
from PIL import Image
from docx import Document
from docx.shared import Mm
from pillow_heif import register_heif_opener
from functions import *

TEMPLATE_PATH = './template/template.docx'
FINAL_PATH = './data/final/'

register_heif_opener()

def create_docx(email_id):
    doc = Document(TEMPLATE_PATH)

    data = {
        'registration_number': regd_from_id(email_id),
        'email': email_id,
        'mobile_number': mobile_from_id(email_id),
        'name': name_from_id(email_id),
        'class_section': class_section_from_id(email_id),
        'previous_positions_held': prev_positions_from_id(email_id),
        'primary_post_applied': post_prim_from_id(email_id),
        'secondary_post_applied': post_sec_from_id(email_id),
        'answer_1': responses_from_id(email_id)[0],
        'answer_2': responses_from_id(email_id)[1],
        'answer_3': responses_from_id(email_id)[2],
    }
    image_path = f'./data/photos/{details_from_id(email_id)[photo_name_num]}'
    print(image_path)
    for table in doc.tables:
        for row in table.rows:
            left_cell, right_cell = row.cells

            # Insert image if placeholder is found in the right cell
            if '{{ photograph }}' in right_cell.text:
                right_cell.text = ""
                if not os.path.exists(image_path):
                    image_path+='.jpg'
                print(image_path)
                with open(image_path, 'rb') as _image:
                    image_bytes = _image.read()
                image = Image.open(io.BytesIO(image_bytes))
                # Get original image dimensions
                width, height = image.size

                # Calculate target aspect ratio (7:9)
                target_ratio = 7 / 9
                # Calculate the new dimensions to maintain 7:9 aspect ratio
                if width / height > target_ratio:  # Image is too wide, crop width
                    new_width = int(height * target_ratio)
                    offset = (width - new_width) // 2  # Center the crop
                    crop_box = (offset, 0, offset + new_width, height)
                else:  # Image is too tall, crop height
                    new_height = int(width / target_ratio)
                    offset = (height - new_height) // 2  # Center the crop
                    crop_box = (0, offset, width, offset + new_height)
                cropped_image = image.crop(crop_box).convert('RGB')
                cropped_image.save(f"./data/photos/cropped/{regd_from_id(email_id)}.jpg")
                right_cell_paragraph = right_cell.paragraphs[0]
                run = right_cell_paragraph.add_run()
                run.add_picture(f"./data/photos/cropped/{regd_from_id(email_id)}.jpg", width=Mm(35), height=Mm(45))

            # Replace placeholders in the left cell text
            for placeholder, value in data.items():
                for paragraph in left_cell.paragraphs:
                    for run in paragraph.runs:
                        if f"{{{{ {placeholder} }}}}" in run.text:
                            # Split text around the placeholder and replace it with value
                            text_parts = run.text.split(f"{{{{ {placeholder} }}}}")
                            run.text = text_parts[0]  # Keep initial part of the run text
                            # Add new run for the placeholder replacement while copying formatting
                            new_run = paragraph.add_run(value)
                            new_run.bold = run.bold
                            new_run.italic = run.italic
                            new_run.font.size = run.font.size
                            # Add rest of text after placeholder
                            if len(text_parts) > 1:
                                end_run = paragraph.add_run(text_parts[1])
                                end_run.bold = run.bold
                                end_run.italic = run.italic
                                end_run.font.size = run.font.size

    # Replace placeholders in document paragraphs (outside the tables)
    for paragraph in doc.paragraphs:
        for placeholder, value in data.items():
            if f"{{{{ {placeholder} }}}}" in paragraph.text:
                for run in paragraph.runs:
                    if f"{{{{ {placeholder} }}}}" in run.text:
                        run.text = run.text.replace(f"{{{{ {placeholder} }}}}", value)

    # Save the document
    doc.save(f"{FINAL_PATH}/{regd_from_id(email_id)} - {name_from_id(email_id)}.docx")


def convert_docx_to_pdf(email_id):
    docx_path = f"{FINAL_PATH}/{regd_from_id(email_id)} - {name_from_id(email_id)}.docx"
    pdf_path = f"{FINAL_PATH}/{regd_from_id(email_id)} - {name_from_id(email_id)}.pdf"

    subprocess.run(
        [r"C:\Program Files\LibreOffice\program\soffice.exe", "--headless", "--convert-to", "pdf", "--outdir",
         os.path.dirname(pdf_path), docx_path])

    if os.path.exists(pdf_path):
        os.remove(docx_path)